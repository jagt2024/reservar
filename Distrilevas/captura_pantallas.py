import streamlit as st
import pyautogui
import time
from PIL import Image
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

# Constantes
SCREENSHOTS_FOLDER = "screenshots"
DOCUMENT_NAME = "registro_capturas.docx"

def create_folders():
    """Crea la carpeta necesaria si no existe"""
    if not os.path.exists(SCREENSHOTS_FOLDER):
        os.makedirs(SCREENSHOTS_FOLDER)

def get_or_create_document():
    """Obtiene el documento existente o crea uno nuevo"""
    if os.path.exists(DOCUMENT_NAME):
        return Document(DOCUMENT_NAME)
    else:
        doc = Document()
        doc.add_heading('Registro de Capturas de Pantalla', 0)
        return doc

def capture_screen():
    """Captura la pantalla y retorna la imagen"""
    countdown = 3
    while countdown > 0:
        st.write(f"Capturando en {countdown}...")
        time.sleep(1)
        countdown -= 1
    
    screenshot = pyautogui.screenshot()
    return screenshot

def save_screenshot(screenshot):
    """Guarda la captura y la añade al documento Word único"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Guardar imagen
    img_filename = f"{SCREENSHOTS_FOLDER}/captura_{timestamp}.png"
    screenshot.save(img_filename)
    
    # Obtener o crear documento Word
    doc = get_or_create_document()
    
    # Añadir la nueva captura al documento
    doc.add_heading(f'Captura de Pantalla', level=1)
    doc.add_paragraph(f'Fecha y hora: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_picture(img_filename, width=Inches(6))
    doc.add_paragraph('-------------------------------------------')
    
    # Guardar documento
    doc.save(DOCUMENT_NAME)
    
    return img_filename

def main():
    st.title("📸 Capturador de Pantallas")
    st.write("Todas las capturas se guardan en un único documento Word")
    
    # Crear carpeta necesaria
    create_folders()
    
    # Mostrar información del documento
    if os.path.exists(DOCUMENT_NAME):
        doc_size = os.path.getsize(DOCUMENT_NAME) / 1024  # Tamaño en KB
        st.info(f"📄 Documento actual: {DOCUMENT_NAME} ({doc_size:.1f} KB)")
    
    # Botón para capturar
    if st.button("📷 Tomar Captura"):
        with st.spinner("Preparando captura..."):
            # Minimizar la ventana de Streamlit
            pyautogui.hotkey('win', 'down')
            time.sleep(1)
            
            # Capturar pantalla
            screenshot = capture_screen()
            
            # Guardar la captura
            img_filename = save_screenshot(screenshot)
            
            # Mostrar información
            st.success("¡Captura guardada exitosamente en el documento!")
            st.image(screenshot, caption="Vista previa de la captura", use_column_width=True)
    
    # Botón para crear nuevo documento
    if st.button("📄 Crear Nuevo Documento"):
        if os.path.exists(DOCUMENT_NAME):
            # Crear backup del documento anterior
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_name = f"registro_capturas_backup_{timestamp}.docx"
            os.rename(DOCUMENT_NAME, backup_name)
            st.success(f"Documento anterior respaldado como: {backup_name}")
        
        # Crear nuevo documento
        doc = Document()
        doc.add_heading('Registro de Capturas de Pantalla', 0)
        doc.save(DOCUMENT_NAME)
        st.success("Nuevo documento creado exitosamente!")

if __name__ == "__main__":
    main()