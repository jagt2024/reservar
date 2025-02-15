import streamlit as st
import pyautogui
import time
import cv2
import numpy as np
from PIL import Image
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import threading
import atexit
import signal
import psutil
import sys
import ctypes
from pathlib import Path

# Constantes
SCREENSHOTS_FOLDER = "screenshots"
RECORDINGS_FOLDER = "recordings"
DOCUMENT_NAME = "registro_capturas.docx"
RECORDING_INTERVAL = 3.0  # Intervalo de grabación en segundos

class ScreenRecorder:
    def __init__(self):
        self.is_recording = False
        self.output_video = None
        self.recording_thread = None
        self.output_path = None
        
    def start_recording(self):
        """Inicia la grabación de pantalla"""
        if self.is_recording:
            return
        
        self.is_recording = True
        screen_size = pyautogui.size()
        fps = 1/RECORDING_INTERVAL  # FPS basado en el intervalo (1/3 fps para 3 segundos)
        
        # Nombre del archivo de salida
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.output_path = os.path.join(RECORDINGS_FOLDER, f'grabacion_{timestamp}.avi')
        
        # Configurar el grabador
        fourcc = cv2.VideoWriter_fourcc(*'XVID')
        self.output_video = cv2.VideoWriter(self.output_path, fourcc, fps, screen_size)
        
        # Iniciar thread de grabación
        self.recording_thread = threading.Thread(target=self._record)
        self.recording_thread.start()
        
    def _record(self):
        """Función principal de grabación"""
        try:
            frames_recorded = 0
            next_capture_time = time.time()
            
            while self.is_recording and self.output_video is not None:
                current_time = time.time()
                
                # Verificar si es tiempo de capturar el siguiente frame
                if current_time >= next_capture_time:
                    # Capturar la pantalla
                    frame = pyautogui.screenshot()
                    # Convertir la imagen a formato opencv
                    frame = cv2.cvtColor(np.array(frame), cv2.COLOR_RGB2BGR)
                    # Escribir el frame
                    self.output_video.write(frame)
                    
                    frames_recorded += 1
                    # Actualizar el tiempo para la siguiente captura
                    next_capture_time = current_time + RECORDING_INTERVAL
                    
                    # Mostrar información de grabación si es posible
                    if frames_recorded % 5 == 0:  # Actualizar cada 5 frames
                        duration = frames_recorded * RECORDING_INTERVAL
                        print(f"Grabando... Duración: {duration:.1f} segundos")
                
                # Pequeña pausa para no sobrecargar el CPU
                time.sleep(0.1)
                
        except Exception as e:
            st.error(f"Error durante la grabación: {str(e)}")
        finally:
            self.stop_recording()

    def stop_recording(self):
        """Detiene la grabación y limpia los recursos"""
        self.is_recording = False
        
        # Esperar a que termine el thread de grabación
        if self.recording_thread and self.recording_thread.is_alive():
            self.recording_thread.join(timeout=1)
        
        # Liberar recursos de video
        if self.output_video is not None:
            self.output_video.release()
            self.output_video = None
        
        return self.output_path

def force_delete_file(file_path):
    """
    Intenta forzar la eliminación de un archivo usando diferentes métodos
    según el sistema operativo.
    """
    try:
        # Primero intentamos el método normal
        os.remove(file_path)
        return True
    except Exception as e:
        try:
            # En Windows, intentamos liberar los manejadores del archivo
            if sys.platform == 'win32':
                # Convertir a Path para manejar mejor las rutas
                path = Path(file_path)
                
                # Intentar cerrar cualquier proceso que esté usando el archivo
                for proc in psutil.process_iter(['pid', 'name', 'open_files']):
                    try:
                        for file in proc.open_files():
                            if Path(file.path) == path:
                                proc.terminate()
                                proc.wait(timeout=1)
                    except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                        continue
                
                # Intentar eliminar usando comandos del sistema
                os.system(f'del /F /Q "{file_path}"')
                
                # Si aún existe, usar handle de Windows
                if os.path.exists(file_path):
                    kernel32 = ctypes.windll.kernel32
                    kernel32.SetFileAttributesW(file_path, 128)  # FILE_ATTRIBUTE_NORMAL
                    os.remove(file_path)
                
            else:  # Para sistemas Unix/Linux
                # Intentar usar comando rm con force
                os.system(f'rm -f "{file_path}"')
            
            return True
        except Exception as inner_e:
            st.error(f"No se pudo eliminar el archivo. Error: {str(inner_e)}")
            return False

def cleanup_recordings():
    """Limpia cualquier grabación que pudiera haber quedado bloqueada"""
    recordings_dir = Path(RECORDINGS_FOLDER)
    if recordings_dir.exists():
        for file in recordings_dir.glob("*.avi"):
            try:
                # Intentar liberar el archivo si está siendo usado por opencv
                cv2.destroyAllWindows()
                # Forzar el cierre de cualquier stream de video
                for obj in gc.get_objects():
                    if isinstance(obj, cv2.VideoWriter):
                        obj.release()
            except:
                pass

def create_folders():
    """Crea las carpetas necesarias si no existen"""
    for folder in [SCREENSHOTS_FOLDER, RECORDINGS_FOLDER]:
        if not os.path.exists(folder):
            os.makedirs(folder)

def get_or_create_document():
    """Obtiene el documento existente o crea uno nuevo"""
    if os.path.exists(DOCUMENT_NAME):
        return Document(DOCUMENT_NAME)
    else:
        doc = Document()
        doc.add_heading('Registro de Capturas de Pantalla', 0)
        return doc

def capture_screen():
    """Captura la pantalla y retorna la imagen"""
    countdown = 3
    while countdown > 0:
        st.write(f"Capturando en {countdown}...")
        time.sleep(1)
        countdown -= 1
    
    screenshot = pyautogui.screenshot()
    return screenshot

def save_screenshot(screenshot):
    """Guarda la captura y la añade al documento Word único"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Guardar imagen
    img_filename = f"{SCREENSHOTS_FOLDER}/captura_{timestamp}.png"
    screenshot.save(img_filename)
    
    # Obtener o crear documento Word
    doc = get_or_create_document()
    
    # Añadir la nueva captura al documento
    doc.add_heading(f'Captura de Pantalla', level=1)
    doc.add_paragraph(f'Fecha y hora: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_picture(img_filename, width=Inches(6))
    doc.add_paragraph('-------------------------------------------')
    
    # Guardar documento
    doc.save(DOCUMENT_NAME)
    
    return img_filename

def get_video_duration(video_path):
    """Calcula la duración aproximada del video basado en su tamaño y FPS"""
    if os.path.exists(video_path):
        size_bytes = os.path.getsize(video_path)
        # Estimación aproximada basada en el tamaño del archivo y el intervalo de grabación
        frames = size_bytes / (1024 * 1024)  # Aproximadamente 1MB por frame
        duration = frames * RECORDING_INTERVAL
        return duration
    return 0

def main():

    cleanup_recordings()
    
    st.title("📸 Capturador y Grabador de Pantalla")
    st.write("Captura imágenes y graba tu pantalla (un frame cada 3 segundos)")
    
    # Crear carpetas necesarias
    create_folders()
    
    # Inicializar el grabador en session_state si no existe
    if 'recorder' not in st.session_state:
        st.session_state.recorder = ScreenRecorder()
    
    # Crear tabs para separar funcionalidades
    tab1, tab2 = st.tabs(["📸 Captura de Pantalla", "🎥 Grabación de Pantalla"])
    
    with tab1:
        # Mostrar información del documento
        if os.path.exists(DOCUMENT_NAME):
            doc_size = os.path.getsize(DOCUMENT_NAME) / 1024  # Tamaño en KB
            st.info(f"📄 Documento actual: {DOCUMENT_NAME} ({doc_size:.1f} KB)")
        
        # Botón para capturar
        if st.button("📷 Tomar Captura"):
            with st.spinner("Preparando captura..."):
                # Minimizar la ventana de Streamlit
                pyautogui.hotkey('win', 'down')
                time.sleep(1)
                
                # Capturar pantalla
                screenshot = capture_screen()
                
                # Guardar la captura
                img_filename = save_screenshot(screenshot)
                
                # Mostrar información
                st.success("¡Captura guardada exitosamente en el documento!")
                st.image(screenshot, caption="Vista previa de la captura", use_column_width=True)
        
        # Botón para crear nuevo documento
        if st.button("📄 Crear Nuevo Documento"):
            if os.path.exists(DOCUMENT_NAME):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup_name = f"registro_capturas_backup_{timestamp}.docx"
                os.rename(DOCUMENT_NAME, backup_name)
                st.success(f"Documento anterior respaldado como: {backup_name}")
            
            doc = Document()
            doc.add_heading('Registro de Capturas de Pantalla', 0)
            doc.save(DOCUMENT_NAME)
            st.success("Nuevo documento creado exitosamente!")
    
    with tab2:
        st.info(f"⏱️ Intervalo de grabación: {RECORDING_INTERVAL} segundos por frame")
        
        # Botones para controlar la grabación
        col1, col2 = st.columns(2)
        
        with col1:
            if not st.session_state.recorder.is_recording:
                if st.button("🎥 Iniciar Grabación"):
                    st.session_state.recorder.start_recording()
                    st.success("¡Grabación iniciada!")
                    st.rerun()
        
        with col2:
            if st.session_state.recorder.is_recording:
                if st.button("⏹️ Detener Grabación"):
                    output_path = st.session_state.recorder.stop_recording()
                    if output_path:
                        st.success(f"¡Grabación guardada en: {output_path}!")
                    st.rerun()
        
        # Mostrar estado actual de la grabación
        if st.session_state.recorder.is_recording:
            st.warning("⚠️ Grabación en curso... (Capturando un frame cada 3 segundos)")
        
        # Mostrar grabaciones existentes
        st.subheader("Grabaciones Guardadas")
        recordings = [f for f in os.listdir(RECORDINGS_FOLDER) if f.endswith('.avi')]
        
        if recordings:
            # Botón para eliminar todas las grabaciones
            if st.button("🗑️ Eliminar todas las grabaciones"):
                try:
                    for recording in recordings:
                        os.remove(os.path.join(RECORDINGS_FOLDER, recording))
                    st.success("¡Todas las grabaciones han sido eliminadas!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al eliminar las grabaciones: {str(e)}")
            
            # Mostrar cada grabación con sus detalles y botón de eliminación
            for recording in recordings:
                recording_path = os.path.join(RECORDINGS_FOLDER, recording)
                size_mb = os.path.getsize(recording_path) / (1024 * 1024)
                duration = get_video_duration(recording_path)
                
                # Usar columns para organizar la información y el botón de eliminar
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.write(f"🎬 {recording}")
                    st.write(f"   📊 Tamaño: {size_mb:.1f} MB")
                    st.write(f"   ⏱️ Duración aproximada: {duration:.1f} segundos")
                
                with col2:
                    # Botón para eliminar grabación individual
                    if st.button("🗑️ Eliminar", key=f"del_{recording}"):
                        try:
                            recording_path = os.path.join(RECORDINGS_FOLDER, recording)
                            # Intentar forzar la eliminación
                            if force_delete_file(recording_path):
                                st.success(f"¡Grabación {recording} eliminada!")
                                time.sleep(1)  # Pequeña pausa para asegurar que el archivo se eliminó
                                st.rerun()
                            else:
                                st.error(f"No se pudo eliminar la grabación {recording}. " 
                                        "Por favor, cierra cualquier programa que pueda estar usando el archivo.")
                        except Exception as e:
                            st.error(f"Error al eliminar la grabación: {str(e)}")

                
                st.write("---")
        else:
            st.info("No hay grabaciones guardadas.")
# Función para limpiar recursos al cerrar
def cleanup():
    if hasattr(st.session_state, 'recorder'):
        st.session_state.recorder.stop_recording()

# Registrar función de limpieza
atexit.register(cleanup)

if __name__ == "__main__":
    main()